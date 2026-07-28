"""Shared agent tools: RAG query, PDF generation, email notification."""
from __future__ import annotations

import base64
import binascii
import io
import re
from datetime import datetime, timedelta, timezone
from typing import Any

import httpx
import structlog

from app.services.llm_client import (
    ILMU_CHAT_MODEL,
    ILMU_EMBEDDING_MODEL,
    ilmu_client,
    openai_client,
    OPENAI_EMBEDDING_MODEL,
)
from app.services.vector_store import ChunkResult, hybrid_search
from core.config import settings

log = structlog.get_logger(__name__)


async def _embed(query: str) -> list[float]:
    try:
        resp = await ilmu_client.embeddings.create(input=query, model=ILMU_EMBEDDING_MODEL)
        return resp.data[0].embedding
    except Exception as exc:
        if openai_client is None:
            raise RuntimeError("No embedding provider available") from exc
        resp = await openai_client.embeddings.create(input=query, model=OPENAI_EMBEDDING_MODEL)
        return resp.data[0].embedding


async def query_rag(
    query: str,
    domain: str,
    *,
    language: str = "bm",
    top_k: int = 5,
) -> list[dict[str, Any]]:
    """Hybrid-search a domain and return serialisable chunk dicts."""
    embedding = await _embed(query)
    chunks: list[ChunkResult] = await hybrid_search(
        query_embedding=embedding,
        query_text=query,
        domain=domain,
        language=language,
        top_k=top_k,
    )
    return [
        {
            "id": c.id,
            "content": c.content,
            "source_title": c.source_title,
            "source_url": c.source_url,
            "ministry": c.ministry,
            "similarity": c.similarity,
        }
        for c in chunks
    ]


def _chunks_to_findings(chunks: list[dict[str, Any]], domain: str) -> list[dict[str, Any]]:
    findings: list[dict[str, Any]] = []
    for c in chunks[:3]:
        findings.append({
            "domain": domain,
            "summary": c.get("content", "")[:400],
            "source_title": c.get("source_title", ""),
            "source_url": c.get("source_url", ""),
            "similarity": c.get("similarity", 0.0),
        })
    return findings


async def query_rag_findings(
    query: str,
    domain: str,
    language: str = "bm",
) -> list[dict[str, Any]]:
    chunks = await query_rag(query, domain, language=language)
    return _chunks_to_findings(chunks, domain)


def extract_pdf_text(document_base64: str) -> str:
    """Extract plain text from a base64-encoded PDF. Falls back to empty string."""
    try:
        raw = base64.b64decode(document_base64, validate=True)
    except (ValueError, binascii.Error):
        return ""
    try:
        from pypdf import PdfReader  # type: ignore[import-untyped]

        reader = PdfReader(io.BytesIO(raw))
        parts = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(parts).strip()
    except Exception as exc:
        log.warning("pdf_extract_failed", error=str(exc))
        return ""


def extract_questions_from_text(text: str, limit: int = 10) -> list[str]:
    """Heuristic question extraction from exam paper text."""
    numbered = re.findall(r"(?:^|\n)\s*(?:\d+[\).:]|Soalan\s+\d+)\s*(.+?)(?=\n\s*(?:\d+[\).:]|Soalan\s+\d+)|\Z)", text, re.DOTALL | re.IGNORECASE)
    cleaned = [re.sub(r"\s+", " ", q).strip() for q in numbered if len(q.strip()) > 10]
    if cleaned:
        return cleaned[:limit]
    lines = [ln.strip() for ln in text.splitlines() if "?" in ln or ln.strip().endswith("?")]
    return lines[:limit]


async def llm_complete(
    system: str,
    user: str,
    *,
    language: str = "bm",
    max_tokens: int = 512,
) -> str:
    """Single-shot ILMU chat completion for agent nodes."""
    lang_note = "Respond in Bahasa Malaysia." if language == "bm" else "Respond in English."
    try:
        resp = await ilmu_client.chat.completions.create(
            model=ILMU_CHAT_MODEL,
            messages=[
                {"role": "system", "content": f"{system}\n{lang_note}"},
                {"role": "user", "content": user},
            ],
            max_tokens=max_tokens,
            temperature=0.2,
        )
        return (resp.choices[0].message.content or "").strip()
    except Exception as exc:
        log.warning("llm_complete_failed", error=str(exc))
        return ""


async def generate_pdf(
    html: str,
    *,
    user_id: str,
    agent_type: str = "compliance-drafter",
    supabase_client: Any = None,
) -> tuple[str, str, str]:
    """Render HTML to PDF, upload to Supabase Storage. Returns (path, signed_url, expires_at)."""
    pdf_bytes: bytes
    try:
        from weasyprint import HTML  # type: ignore[import-untyped]

        pdf_bytes = HTML(string=html).write_pdf()
    except Exception as exc:
        log.warning("weasyprint_unavailable", error=str(exc))
        pdf_bytes = html.encode("utf-8")

    storage_path = f"agents/{agent_type}/{user_id}/{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}.pdf"
    bucket = settings.supabase_storage_bucket

    if not supabase_client:
        return storage_path, "", ""

    try:
        supabase_client.storage.from_(bucket).upload(
            storage_path,
            pdf_bytes,
            {"content-type": "application/pdf", "upsert": "true"},
        )
        expires_at = datetime.now(timezone.utc) + timedelta(hours=24)
        signed = supabase_client.storage.from_(bucket).create_signed_url(
            storage_path,
            86_400,
        )
        url = signed.get("signedURL") or signed.get("signedUrl") or ""
        return storage_path, url, expires_at.isoformat()
    except Exception as exc:
        log.warning("pdf_upload_failed", error=str(exc))
        return storage_path, "", ""


async def generate_docx(
    report_json: dict[str, Any],
    *,
    user_id: str,
    agent_type: str = "grant-draft-generator",
    supabase_client: Any = None,
) -> tuple[str, str, str]:
    """Render a structured report to a real Word document (headings per
    section, not HTML dumped into a text run), upload to Supabase Storage.
    Returns (path, signed_url, expires_at). Same storage/signed-URL pattern
    as generate_pdf, namespaced by agent_type."""
    docx_bytes: bytes
    try:
        from docx import Document  # type: ignore[import-untyped]

        doc = Document()
        doc.add_heading(
            f"Grant Application Draft — {report_json.get('programme_name', '')}", level=0
        )
        disclaimer = report_json.get("disclaimer", "")
        if disclaimer:
            p = doc.add_paragraph()
            run = p.add_run(disclaimer)
            run.italic = True

        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(report_json.get("executive_summary", ""))

        doc.add_heading("Use of Funds", level=1)
        doc.add_paragraph(report_json.get("use_of_funds_narrative", ""))

        doc.add_heading("Financial Projection Skeleton", level=1)
        skeleton = report_json.get("financial_projection_skeleton") or {}
        skeleton_disclaimer = skeleton.get("disclaimer", "")
        if skeleton_disclaimer:
            p = doc.add_paragraph()
            run = p.add_run(skeleton_disclaimer)
            run.italic = True
            run.bold = True

        doc.add_heading("Revenue Projection (template)", level=2)
        for row in skeleton.get("revenue_projection", []):
            doc.add_paragraph(f"{row.get('period')}: RM_____ ({row.get('notes', '')})", style="List Bullet")

        doc.add_heading("Cost Breakdown (template)", level=2)
        for row in skeleton.get("cost_breakdown", []):
            doc.add_paragraph(f"{row.get('category')}: RM_____", style="List Bullet")

        doc.add_heading("Funding Allocation (template)", level=2)
        for row in skeleton.get("funding_allocation", []):
            doc.add_paragraph(f"{row.get('use_of_funds')}: RM_____", style="List Bullet")

        doc.add_heading("Required Document Checklist", level=1)
        for item in report_json.get("document_checklist", []) or []:
            required = "Required" if item.get("required") else "Optional"
            doc.add_paragraph(
                f"{item.get('item')} ({required}): {item.get('description')}",
                style="List Bullet",
            )

        if disclaimer:
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(disclaimer)
            run.italic = True

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()
    except Exception as exc:
        log.warning("python_docx_unavailable", error=str(exc))
        docx_bytes = str(report_json).encode("utf-8")

    storage_path = f"agents/{agent_type}/{user_id}/{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}.docx"
    bucket = settings.supabase_storage_bucket

    if not supabase_client:
        return storage_path, "", ""

    try:
        supabase_client.storage.from_(bucket).upload(
            storage_path,
            docx_bytes,
            {
                "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "upsert": "true",
            },
        )
        expires_at = datetime.now(timezone.utc) + timedelta(hours=24)
        signed = supabase_client.storage.from_(bucket).create_signed_url(
            storage_path,
            86_400,
        )
        url = signed.get("signedURL") or signed.get("signedUrl") or ""
        return storage_path, url, expires_at.isoformat()
    except Exception as exc:
        log.warning("docx_upload_failed", error=str(exc))
        return storage_path, "", ""


async def send_email(
    *,
    to: str,
    subject: str,
    html_body: str,
) -> bool:
    """Send via Resend API. Returns False when not configured."""
    api_key = settings.resend_api_key.strip()
    from_addr = settings.resend_from_email.strip()
    if not api_key or not to:
        log.info("send_email_skipped", reason="not_configured")
        return False

    async with httpx.AsyncClient(timeout=15.0) as client:
        resp = await client.post(
            "https://api.resend.com/emails",
            headers={"Authorization": f"Bearer {api_key}"},
            json={"from": from_addr, "to": [to], "subject": subject, "html": html_body},
        )
        if resp.status_code >= 400:
            log.warning("send_email_failed", status=resp.status_code)
            return False
    return True


async def grant_compatibility_check(
    programme_names: list[str],
    supabase: Any,
    *,
    language: str = "en",
) -> dict[str, Any]:
    """Grant stacking compatibility matrix (stackable / partial_overlap /
    conflict / unknown) for a set of programmes a founder plans to apply for
    simultaneously.

    Thin delegation to app.agents.eligibility_agent.compatibility so this
    module stays the single import surface for agent-callable tools. Imported
    lazily to keep tools.py free of an eligibility-agent import cycle.
    """
    from app.agents.eligibility_agent.compatibility import (
        grant_compatibility_check as _check,
    )

    return await _check(programme_names, supabase, language=language)
